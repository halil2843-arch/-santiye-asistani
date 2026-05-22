"""Şablon alan tespiti servisi.

Excel (.xlsx) ve Word (.docx) dosyalarındaki {{ alan_adi }} formatındaki
yer tutucuları bulur ve yapılandırılmış biçimde döndürür.
"""

import re
from pathlib import Path

import openpyxl
from docx import Document


_PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*(\w+)\s*\}\}")


def parse_xlsx_fields(dosya_yolu: str) -> dict[str, str]:
    """Excel dosyasındaki {{ alan_adi }} yer tutucularını hücre koordinatlarıyla döndürür.

    Args:
        dosya_yolu: .xlsx dosyasının tam yolu.

    Returns:
        Hücre koordinatı → alan adı eşlemesi.
        Örnek: {"B3": "tarih", "D8": "proje_muduru"}

    Raises:
        FileNotFoundError: Dosya bulunamazsa.
        ValueError: Geçersiz Excel dosyası ise.
    """
    path = Path(dosya_yolu)
    if not path.exists():
        raise FileNotFoundError(f"Excel dosyası bulunamadı: {dosya_yolu}")

    result: dict[str, str] = {}

    try:
        wb = openpyxl.load_workbook(path, data_only=False)
    except Exception as exc:
        raise ValueError(f"Excel dosyası açılamadı: {exc}") from exc

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    matches = _PLACEHOLDER_PATTERN.findall(cell.value)
                    for alan_adi in matches:
                        coord = cell.coordinate  # örn. "B3"
                        result[coord] = alan_adi

    wb.close()
    return result


def parse_docx_fields(dosya_yolu: str) -> list[str]:
    """Word dosyasındaki {{ alan_adi }} yer tutucularını listeler.

    Paragraflar, tablolar ve header/footer alanlarını tarar.

    Args:
        dosya_yolu: .docx dosyasının tam yolu.

    Returns:
        Tekrarsız alan adları listesi. Örnek: ["tarih", "hava_durumu"]

    Raises:
        FileNotFoundError: Dosya bulunamazsa.
        ValueError: Geçersiz Word dosyası ise.
    """
    path = Path(dosya_yolu)
    if not path.exists():
        raise FileNotFoundError(f"Word dosyası bulunamadı: {dosya_yolu}")

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"Word dosyası açılamadı: {exc}") from exc

    found: list[str] = []

    # Paragraflar
    for para in doc.paragraphs:
        found.extend(_PLACEHOLDER_PATTERN.findall(para.text))

    # Tablolar
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    found.extend(_PLACEHOLDER_PATTERN.findall(para.text))

    # Header ve footer bölümleri
    for section in doc.sections:
        for hdr_para in section.header.paragraphs:
            found.extend(_PLACEHOLDER_PATTERN.findall(hdr_para.text))
        for ftr_para in section.footer.paragraphs:
            found.extend(_PLACEHOLDER_PATTERN.findall(ftr_para.text))

    # Tekrarları kaldır, sırayı koru
    seen: set[str] = set()
    unique: list[str] = []
    for alan in found:
        if alan not in seen:
            seen.add(alan)
            unique.append(alan)

    return unique
